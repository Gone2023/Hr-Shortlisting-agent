import fitz  # PyMuPDF
from docx import Document
import os

def parse_pdf(file_path: str) -> str:
    """Extracts text from a PDF file."""
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text() + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def parse_docx(file_path: str) -> str:
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def parse_document(file_path: str) -> str:
    """Parses document based on extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext == '.docx':
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

def parse_document_bytes(file_bytes: bytes, filename: str) -> str:
    """Parses document from bytes (useful for Streamlit uploads)."""
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    if ext == '.pdf':
        try:
            with fitz.open("pdf", file_bytes) as doc:
                for page in doc:
                    text += page.get_text() + "\n"
        except Exception as e:
            print(f"Error reading PDF bytes: {e}")
    elif ext == '.docx':
        import io
        try:
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX bytes: {e}")
    elif ext == '.json':
        import json
        try:
            data = json.loads(file_bytes.decode('utf-8'))
            text = json.dumps(data, indent=2)
        except Exception as e:
            print(f"Error reading JSON bytes: {e}")
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
    return text
